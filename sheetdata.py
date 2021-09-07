# import smtplib
# from email.mime.multipart import MIMEMultipart
# from email.mime.text import MIMEText
# from email.mime.application import MIMEApplication
import datetime
import errno
import json
import time
import gspread
import comtypes.client
import autoit
import openpyxl
import pdfrw
from pdfrw import PdfWriter
from oauth2client.service_account import ServiceAccountCredentials
from openpyxl import Workbook
import os
from google_drive_downloader import GoogleDriveDownloader as gdd
from collections import OrderedDict
from docx import Document
import math
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains


def get_records():
    scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
    creds = ServiceAccountCredentials.from_json_keyfile_name('HR Onboarding-2eaf151a35d8.json', scope)
    client = gspread.authorize(creds)
    datasheet = client.open("HR ROBO SHEET").sheet1
    sheet2 = client.open("HR ROBO SHEET").worksheet('Sheet2')
    dict_records = datasheet.get_all_records()
    list_records = datasheet.get_all_values()
    code_records = OrderedDict(sheet2.get_all_records()[0])
    return dict_records, list_records, code_records


def create_json(path, key, value):
    try:
        with open(os.path.join(path, "emaildates.json")) as json_file:
            data = json.loads(json.load(json_file))
            data[key] = value
            data = json.dumps(data)
    except:
        data = {}
        data[key] = value
        data = json.dumps(data)

    with open(os.path.join(path, "emaildates.json"), 'w+') as json_data:
        json.dump(data, json_data)


def sendoutlookmail(to, cc, subj, body, att):
    driver = webdriver.Chrome("chromedriver.exe")
    driver.maximize_window()
    driver.get("https://webmail.mahindra.com/owa/#path=/mail")
    WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.ID, 'username')))
    driver.find_element_by_id("username").send_keys("mahindra\onboardingteam-mibs")
    driver.find_element_by_id("password").send_keys("mahindra@123")
    driver.find_element_by_class_name("signinbutton").click()
    WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.ID, '_ariaId_39')))
    driver.find_element_by_id("_ariaId_39").click()
    WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.CLASS_NAME, '_fp_u')))
    st = ""
    for item in att:
        st += F'"{item}" '
    driver.find_element_by_xpath(
        '/html/body/div[2]/div/div[3]/div[3]/div/div[1]/div[2]/div[7]/div/div/div[2]/div[1]/span[1]/div[3]/button').click()
    WebDriverWait(driver, 20).until(EC.element_to_be_clickable((By.CLASS_NAME, '_fce_b')))
    driver.find_element_by_class_name('_fce_b').click()
    time.sleep(1)
    autoit.control_focus("Open", "Edit1")
    autoit.control_set_text("Open", "Edit1", st)
    autoit.control_click("Open", "Button1")
    time.sleep(2)

    to_ele = driver.find_element_by_xpath(
        '/html/body/div[2]/div/div[3]/div[3]/div/div[1]/div[2]/div[7]/div/div/div[3]/div[2]/div[1]/div[3]/div[1]/div[2]/div/div/span/div[1]/form/input')
    time.sleep(2)
    to_ele.click()
    for x in to:
        to_ele.send_keys(x)
        time.sleep(1)
        to_ele.send_keys("\t")
    time.sleep(0.5)

    to_ele = driver.find_element_by_xpath(
        '/html/body/div[2]/div/div[3]/div[3]/div/div[1]/div[2]/div[7]/div/div/div[3]/div[2]/div[1]/div[4]/div[1]/div[2]/div/div/span/div[1]/form/input')
    to_ele.click()
    for x in cc:
        to_ele.send_keys(x)
        time.sleep(1)
        to_ele.send_keys("\t")
    time.sleep(0.5)
    ActionChains(driver).key_down(Keys.TAB).key_up(Keys.TAB).perform()

    driver.find_element_by_xpath(
        '/html/body/div[2]/div/div[3]/div[3]/div/div[1]/div[2]/div[7]/div/div/div[3]/div[2]/div[1]/div[7]/div/div/input').send_keys(
        subj)
    ActionChains(driver).key_down(Keys.TAB).key_down(Keys.TAB).key_down(Keys.TAB).key_down(Keys.TAB).key_down(
        Keys.TAB).key_down(Keys.TAB).key_down(Keys.TAB).key_down(Keys.TAB).send_keys(body).perform()
    time.sleep(1)
    driver.find_element_by_class_name("_mcp_M1").click()
    time.sleep(5)


def create_id_staffing(records, fpath):
    book = Workbook()
    sheet = book.active
    sheet.append(["FIRST NAME", "LAST NAME", "EMPLOYER NUMBER (NEW)", "DESIGNATION", "LOCATION",
                  "IF TRAINEE OR CONTRACT DD/MM/YYYY", "LAST DATE OF CONTRACT", "IMMEDIATE SUPERIOR Name & Token ID  ",
                  "(MAILBOX ID ONLY)", "Mobile", "Home Phone", "EXTN", "Sectore Code(Mandatory)",
                  "O365 license (E1 or E3)", "Company code(Mandatory)", "Costcenter id(Mandatory)"])
    sheet.append(
        [records[2], records[4], records[57], records[67], records[70], "", "", records[76] + " " + records[77],
         records[91], records[36], records[31], "", "MIBS", "", "MB10381221", records[68]])
    book.save(fpath)


def create_id_core(records, fpath):
    book = Workbook()
    sheet = book.active
    sheet.append(
        ["FIRST NAME", "LAST NAME", "EMPLOYER NUMBER (6 DIGIT NUMARIC ONLY)", "DESIGNATION", "LOCATION", "Department",
         "IF TRAINEE OR CONTRACT DD/MM/YYYY", "LAST DATE OF CONTRACT.", "IMMEDIATE SUPERIOR Name & Token ID",
         "(MAILBOX ID ONLY)", "Mobile", "EXTN", "Sectore Code(Mandatory)", "O365 license (E1 or E3)",
         "Company code(Mandatory)", "Costcenter id(Mandatory)", "Address", "Pin code"])

    sheet.append([records[2], records[4], records[57], records[67], records[70], records[84], "", "",
                  records[76] + " " + records[77], records[91], records[36], "", "MIBS", "", "1038", "MAS9905", "", ""])
    book.save(fpath)


def create_id_card(records, fpath, loc1):
    doj = records[58]
    doj = doj[0:2] + '/' + doj[3:5] + '/' + doj[-4:]
    book = Workbook()
    sheet = book.active
    sheet.append(["Emp. No", "Employee Name", "Blood Group", "Date of Issue/DOJ", "BAND COLOR", "Location",
                  "EMERGENCY CONTACT NO"])
    sheet.append([records[57], records[8], records[30], doj, "BLUE", loc1, records[31]])
    book.save(fpath)


def create_ascent(records, fpath):
    if not os.path.isfile(fpath):
        book = Workbook()
        sheet = book.active
        sheet.append(['Employee ID', 'Gender -M/F', 'First Name', 'Middle Name', 'Last Name', 'Fathers/Husband Name',
                      'EmpRelation', 'Display Name', 'Marital Status', 'Spouse Name', 'No. of Children',
                      'Present Address 1', 'Present Address 2', 'Present Address 3', 'Present City', 'Present State',
                      'Present PinCode', 'Present Phone', 'Permanent Address 1', 'Permanent Address 2',
                      'Permanent Address 3', 'Permanent City', 'Permanent State', 'Permanent PinCode',
                      'Permanent Phone', 'Primary Bank  Name', 'Primary IFSC', 'Primary Bank A/c No', 'Payroll Code',
                      'Date Of Joining', 'Training From', 'Date of Probation', 'Date of Confirmation',
                      'Date of Retirement', 'Date Of Birth', 'Category Code', 'Status Code', 'Grade Code',
                      'Designation', 'Cost Centre Code', 'Business Area Code', 'Location Code', 'Occupation Code',
                      'Qualification', 'Permanent A/c No.', 'P.F. Registration Code', 'P.F. A/c No.(6 Digits)',
                      'PF wef Dt.', 'E.S.I. No.', 'Blood Group', 'Emergency Phone No.', 'Emergency Contact Person',
                      'Email ID', 'Reports To Emp', 'Mobile No ( Reporting manager)', 'Web User Name',
                      'Web User Password', 'Web Access Level', 'Leave Policy ID', 'MICR', 'Attendance User ID',
                      'Designation Code', 'Department Code', 'CR Mem.No.', 'Client Name Code', 'State Code',
                      'Employee Cost center', 'Aadhaar Card No', 'Primary NameAsPerBank', 'UAN', 'Personal Email ID',
                      'Group Joining Date', 'Personal Mobile No'])
    else:
        book = openpyxl.load_workbook(fpath)
        sheet = book.active

    gender = records[5][0]
    father_husband = records[7][0]
    marital_status = ""
    if records[9] == 'Unmarried' and records[5] == 'Male':
        marital_status = 'B'
    elif records[9] == 'Unmarried' and records[5] == 'Female':
        marital_status = 'S'
    elif records[9] == 'Married':
        marital_status = 'M'

    doj = records[58]
    doj = doj[-4:] + '-' + doj[0:2] + '-' + doj[3:5]
    doc = records[62]
    doc = doc[-4:] + '-' + doc[0:2] + '-' + doc[3:5]
    if len(records[28]) == 9:
        dob = '0' + records[28]
        dob = dob[-4:] + '-' + dob[0:2] + '-' + dob[3:5]
    else:
        dob = records[28]
        dob = dob[-4:] + '-' + dob[0:2] + '-' + dob[3:5]

    pf_wef_dt = records[74]
    pf_wef_dt = pf_wef_dt[-4:] + '-' + pf_wef_dt[0:2] + '-' + pf_wef_dt[3:5]

    sheet.append([records[57], gender, records[2], records[3], records[4], records[6], father_husband,
                  records[2] + ' ' + records[3] + ' ' + records[4], marital_status, records[10], records[11],
                  records[12], '', '', records[13], records[14], records[15], records[16], records[12], '', '',
                  records[13], records[14], records[15], records[16], records[22], records[23], records[24],
                  records[63], doj, records[59], records[60], doc, records[61], dob, records[64], records[65],
                  records[66], records[67], records[68], records[69], records[70], records[72], records[29],
                  records[34], '', '', pf_wef_dt, records[75], records[30], records[31], records[32], '', records[76],
                  records[78], records[79], records[80], records[81], '', records[26], records[82], records[83],
                  records[84], records[86], records[88], records[89], '', records[33], records[25], records[46],
                  records[35], '', records[16]])
    book.save(fpath)


def create_zing(records, fpath):
    if not os.path.isfile(fpath):
        book = Workbook()
        sheet = book.active
        sheet.append(["EmpCode", "Salutation", "FirstName", "MiddleName", "LastName", "FatherName", "DateOfBirth",
                      "DateOfJoining", "DateOfConfirmation", "GroupDOJ", "RetirementAge", "EmployeeStatus",
                      "EmployeeGroup", "AttendanceGroup", "LeaveGroup", "CalendarGroup", "AttendanceModeGroup",
                      "AttendanceRuleGroup", "Company", "CostCentre", "Department", "SubDepartment", "Sector", "Client",
                      "ClientLocation", "Category", "Grade", "Designation", "EmployeeLocation", "EmployeesubLocation",
                      "City", "State", "Premises", "EmployeeActivity", "ClaimEligibility", "Group", "PersonnelArea",
                      "Band", "CostCenter", "OrganizationUnitCode", "SubDivision", "PayrollAreaNo.", "BusinessFunction",
                      "BusinessDomain", "DomainType", "CompanyCode", "BusinessUnit", "Division", "CostCenterCode",
                      "OrganizationUnitName", "SubCategory", "PayrollArea", "ProfessionTaxState", "Recruitment",
                      "AttendanceRostering", "StaffingContractSBUNONMSRTEL", "PMS", "CompetencyGroup", "AdminChecklist",
                      "ITChecklist", "Location", "Area", "Zone", "CTCType", "BALFLEXI", "BASIC", "CONTRIBUTIONS",
                      "DEDUCTIONS", "NTH", "HRA", "ProfessionTax", "PARENTSMEDICLAIMDEDUCTION", "Gender", "PANNo",
                      "AadharNo", "UANNo", "OfficialEmailAddress", "Nation", "PresentAddress", "PresentCity",
                      "PermanentAddress", "PermanentState", "PermanentCity", "PresentMobilenumber", "PaymentMode",
                      "BankName", "BankAccountNo", "IFSCCode", "BankBranch", "MaritalStatus", "SpouseName",
                      "DateOfMarriage", "BloodGroup", "ReportingManager", "LMSReportingManager1",
                      "LMSReportingManager2", "LMSReportingManager3", "TNAReportingManager", "WagesCode",
                      "PersonalEmailAddress", "COMPESIC", "COMPPF", "EMPESIC", "EMPPF", "GRATUITY", "PFApplicable",
                      "PFAccountNo", "PFType", "PFBaseSalaryLimit", "PFDenotion", "PFDenotionFigure", "EPSApplicable",
                      "EPSType", "EPSBaseSalaryLimit", "EPSDenotion", "EPSDenotionFigure", "VPFType",
                      "VPFBaseSalaryLimit", "VPFDenotion", "VPFDenotionFigure", "ESICApplicable", "ESICAccountNo",
                      "ESICDenotion", "ESICDenotionFigure", "PTApplicable", "PTState", "LwfApplicable",
                      "Physical_Status", "Employment_Type", "DateOfResignation", "DateOfLeaving",
                      "DateOfLetterSubmission", "Reason", "RemarksForFnF", "CardCode", "Day1", "Frequency1", "Day2",
                      "Frequency2", "AttendanceMode", "Old_EmployeeCode", "ShiftName", "ApprovedRequired"])
    else:
        book = openpyxl.load_workbook(fpath)
        sheet = book.active

    sheet.append(
        [records[57], records[1], records[2], records[3], records[4], records[6], records[28], records[58], "1/1/1990",
         records[90], "", "NewJoinee", "common", "common", "common", "common", "common", "common", "MIBS Staffing",
         "common", "common", "common", "common", records[87], records[70], "Staffing-Contract", "CONTRACT", records[67],
         "common", "common", "", "", "", "", "common", "common", "", "common", "common", "common", "common", "", "", "",
         "", "common", "", "", "common", "common", "common", "common", "Maharashtra", "", "", "", "", "", "", "", "",
         "", "", "", "", "", "", "", "", "", "", "", records[5], records[34], records[33], "", "EMAILID", "INDIA", "",
         "", "", "", "", records[36], "", "", "", "", "", records[9], "", "", "", records[76], "", "", "", "",
         "Regular", "", "", "", "", "", "", "1", "NOPF", "RPFC", "15000", "R", "12", "0", "RPFC", "15000", "R", "8.33",
         "RPFC", "0", "R", "12", "1", "NOESIC", "R", "1.75", "1", "Maharashtra", "1", "Normal", "Full Time", "", "", "",
         "", "", "", "", "", "", "", "", "", "", ""])
    book.save(fpath)


def full_record(records, fpath):
    if not os.path.isfile(fpath):
        book = Workbook()
        sheet = book.active
        sheet.append(["Date and Time", "Salutation/Title", "First Name", "Middle Name", "Last Name", "Gender",
                      "Father’s/Husband Name", "Employee Relation", "Employee Name", "Marital Status", "Spouse Name",
                      "No. of Children", "Present Address", "Present City", "Present State", "Present Pin code",
                      "Present Phone", "Permanent Address", "Permanent City", "Permanent State", "Permanent Pin code",
                      "Permanent Phone", "Primary Bank Name", "Primary IFSC Code", "Primary Account No",
                      "Primary name as per bank", "MICR", "Expected Date of Joining(MM/DD/YYYY)",
                      "Date of Birth(MM/DD/YYYY)", "Qualification", "Blood Group", "Emergency Phone No",
                      "Emergency Contact Person", "Aadhar card no", "Permanent Account No", "Email Address",
                      "Personal Mobile No", "Establishment Address", "Universal Account Number", "PF Account Number",
                      "Date of joining (Unexempted)", "Date of exit (Unexempted)", "Scheme Certificate (if issued)",
                      "PPO Number (if issued)", "Non-Contributory Period (NCP) Days", "Name & Address of the Trust",
                      "UAN", "Member EPS A/c No", "Date of joining (Exempted)", "Date of exit (Exempted)",
                      "Scheme Certificate No (if issued)", "Non Contributory period (NCP) Days",
                      "State Country of origin", "Passport No.", "Validity of Passport", "Upload Passport Photo",
                      "Upload Docs", "Employee Number", "Confirm Date of Joining(MM/DD/YYYY)", "Training From",
                      "Date of Probation(MM/DD/YYYY)", "Date of Retirement(MM/DD/YYYY)",
                      "Date of confirmation(MM/DD/YYYY)", "Payroll Code", "Category Code", "Status Code", "Grade Code",
                      "Designation", "Cost Center Code", "Business Area Code", "Location Code", "Sub-Location Code",
                      "Occupation Code", "PF Registration Code", "PF wef Dt", "E.S.I. No.", "Reports to emp",
                      "Reporting Manager's Token ID", "Mobile number", "Web user name", "Web user password",
                      "Web access level", "Attendance user id", "Designation Code", "Department Code",
                      "Sub-Department Code", "CR Mem No", "Client name", "Client name code", "State code",
                      "Group joining date(MM/DD/YYYY)", "Reporting Manager Email id"])
    else:
        book = openpyxl.load_workbook(fpath)
        sheet = book.active
    rec = []
    rec.append(get_date())
    rec.extend(records[1:])
    sheet.append(rec)
    book.save(fpath)


def create_mis(fpath, empno):
    if not os.path.isfile(fpath):
        book = Workbook()
        sheet = book.active
        sheet.append(["Employee No.", "Status", "Sent On"])
    else:
        book = openpyxl.load_workbook(fpath)
        sheet = book.active
    rec = [empno, "Mail Sent", get_date()]
    sheet.append(rec)
    book.save(fpath)


#
# def send_mails(exl, file, photo=False):
#     fromaddr = "mmtsequel@gmail.com"
#     toaddr = ["danish.khan@sequelstring.com", "samit.pawar@sequelstring.com"]
#     msg = MIMEMultipart()
#     msg['From'] = fromaddr
#     msg['To'] = ", ".join(toaddr)
#     msg['Subject'] = "{} Excel".format(exl)
#     body = """Hello there,
# 		Please Find the attachment for {} Excel""".format(exl)
#
#     msg.attach(MIMEText(body, 'plain'))
#     if photo:
#         attachment = MIMEApplication(open(photo, "rb").read(), _subtype="txt")
#         attachment.add_header('Content-Disposition', 'attachment', filename=os.path.basename(photo))
#         msg.attach(attachment)
#
#     attachment = MIMEApplication(open(os.path.join(exl, file), "rb").read(), _subtype="txt")
#     attachment.add_header('Content-Disposition', 'attachment', filename=os.path.basename(file))
#     msg.attach(attachment)
#     s = smtplib.SMTP('smtp.gmail.com', 587)
#     s.starttls()
#     s.login(fromaddr, "mmt@12345")
#     text = msg.as_string()
#     s.sendmail(fromaddr, toaddr, text)
#     s.quit()

#
# def send_joining_confirmation(rec, toaddr):
#     fromaddr = "mmtsequel@gmail.com"
#     # toaddr = ["danish.khan@sequelstring.com", "samit.pawar@sequelstring.com"]
#     msg = MIMEMultipart()
#     msg['From'] = fromaddr
#     msg['To'] = ", ".join(toaddr)
#     msg['Subject'] = "Appointment Letter"
#     body = f"""
#
# Dear Sir/ Madam,
#
#
#
# We take immense pleasure to introduce (Name- {rec[0]} {rec[1]}), who has joined us as {rec[2]}. She/He will be operating from {rec[4]}.
#
#
#
# Details are as follows
#
# Name: {rec[0]}
#
# Designation: {rec[2]}
#
# Department: {rec[3]}
#
# Location: {rec[4]}
#
# Date: {rec[5]}
#
#
#
#
# Regards,
#
# HR Team."""
#
#     msg.attach(MIMEText(body, 'plain'))
#     s = smtplib.SMTP('smtp.gmail.com', 587)
#     s.starttls()
#     s.login(fromaddr, "mmt@12345")
#     text = msg.as_string()
#     s.sendmail(fromaddr, toaddr, text)
#     s.quit()
#

def get_date():
    dt = datetime.datetime.now()
    date = dt.strftime("%d-%m-%Y")
    return str(date)


def make_dir(*paths):
    # Creates all required directories if not present, mentioned in the path.
    for pt in paths:
        if not (os.path.isdir(pt)):
            try:
                os.makedirs(pt, mode=0o777, exist_ok=True)
            except OSError as exception:
                if exception.errno != errno.EEXIST:
                    raise


def delete_record(record):
    # use creds to create a client to interact with the Google Drive API
    scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
    # creds = ServiceAccountCredentials.from_json_keyfile_name('client_secret.json', scope)
    creds = ServiceAccountCredentials.from_json_keyfile_name('HR Onboarding-2eaf151a35d8.json', scope)
    client = gspread.authorize(creds)

    # Find a workbook by name and open the first sheet
    # Make sure you use the right name here.
    sheet = client.open("HR ROBO SHEET").sheet1
    # sheet.update_cells()
    # Extract and print all of the values
    dict_records = sheet.get_all_records()
    list_records = sheet.get_all_values()
    for i, x in enumerate(list_records):
        if record[2] == x[2] and record[16] == x[16]:
            sheet.delete_row(i + 1)


def save_data(record, index):
    try:
        # use creds to create a client to interact with the Google Drive API
        scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
        # creds = ServiceAccountCredentials.from_json_keyfile_name('client_secret.json', scope)
        creds = ServiceAccountCredentials.from_json_keyfile_name('HR Onboarding-2eaf151a35d8.json', scope)
        client = gspread.authorize(creds)

        # Find a workbook by name and open the first sheet
        # Make sure you use the right name here.
        sheet = client.open("HR ROBO SHEET").sheet1
        # sheet.update_cells()
        # Extract and print all of the values
        cell_list = sheet.range(F'A{index}:CQ{index}')

        for ind, cell in enumerate(cell_list):
            cell.value = record[ind]
        # Update in batch
        sheet.update_cells(cell_list)
    # sheet.insert_row(record,index)
    except Exception as e:
        return "not saved"
    return "saved"


def create_epf(data_dict, template_path, save_path):
    ANNOT_KEY = '/Annots'
    ANNOT_FIELD_KEY = '/T'
    ANNOT_VAL_KEY = '/V'
    ANNOT_RECT_KEY = '/Rect'
    SUBTYPE_KEY = '/Subtype'
    WIDGET_SUBTYPE_KEY = '/Widget'
    template_pdf = pdfrw.PdfReader(template_path)
    template_pdf.Root.AcroForm.update(pdfrw.PdfDict(NeedAppearances=pdfrw.PdfObject('true')))
    pdfObj = PdfWriter()
    annotations = template_pdf.pages[0][ANNOT_KEY]
    for annotation in annotations:
        if annotation[SUBTYPE_KEY] == WIDGET_SUBTYPE_KEY:
            if annotation[ANNOT_FIELD_KEY]:
                key = annotation[ANNOT_FIELD_KEY][1:-1]
                if key in data_dict.keys():
                    annotation.update(pdfrw.PdfDict(V='{}'.format(data_dict[key])))
                    try:
                        pdfObj.write(save_path, template_pdf)
                    except:
                        print("File I/O Exception")


def update_series_data(record, index=2):
    try:
        # use creds to create a client to interact with the Google Drive API
        scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
        # creds = ServiceAccountCredentials.from_json_keyfile_name('client_secret.json', scope)
        creds = ServiceAccountCredentials.from_json_keyfile_name('HR Onboarding-2eaf151a35d8.json', scope)
        client = gspread.authorize(creds)

        # Find a workbook by name and open the first sheet
        # Make sure you use the right name here.
        sheet = client.open("HR ROBO SHEET").worksheet('Sheet2')
        # sheet.update_cells()
        # Extract and print all of the values
        cell_list = sheet.range(F'A{index}:O{index}')

        for ind, cell in enumerate(cell_list):
            cell.value = record[ind]
        # Update in batch
        sheet.update_cells(cell_list)
        # sheet.insert_row(record,index)
        print("Data updated")
    except Exception as e:
        print(e)
        return "not saved"
    return "saved"


def convert_to_pdf(in_file, out_file):
    wdFormatPDF = 17
    # create COM object
    word = comtypes.client.CreateObject('Word.Application')
    # key point 1: make word visible before open a new document
    # word.Visible = True
    # key point 2: wait for the COM Server to prepare well.
    time.sleep(3)
    word.Visible = False
    doc = word.Documents.Open(in_file)  # open docx file 1
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)  # conversion
    doc.Close()


def send_zing(driver_path):
    browser = webdriver.Chrome(driver_path)
    browser.maximize_window()
    browser.get('https://portal.zinghr.com/2015/pages/authentication/MIBS.htm')

    username_field = browser.find_element_by_id('txtEmpCode')
    username_field.clear()

    username_field.send_keys('HR-ROBO')

    password_field = browser.find_element_by_id('txtPassword')
    password_field.clear()

    password_field.send_keys('HR-ROBO123')

    password_login_button = browser.find_element_by_id('btnLogin')
    password_login_button.click()

    time.sleep(3)

    browser.find_element_by_id("sidebar-collapse").click()
    browser.find_element_by_class_name("Payroll").click()
    # browser.get("https://portal.zinghr.com/2015/Pages/Authentication/EmployeeCreationLogin.aspx?Menu=Payroll")
    # time.sleep(111)
    # """
    # menu_button = browser.find_element_by_xpath('//*[@id="sidebar-collapse"]/i')
    # menu_button.click()
    #
    # payroll_button = browser.find_element_by_xpath('//*[@id="menu"]/ul/li[16]/a/i')
    # payroll_button.click()
    #
    # time.sleep(30)
    # browser.find_element_by_tag_name('body').send_keys(Keys.CONTROL + Keys.TAB)
    browser.switch_to.window(browser.window_handles[1])
    # browser.get("https://portal.zinghr.com/NewPayroll/Payroll/PayrollCockpit#NewJoinee-config")
    # browser.find_element_by_class_name("").click()
    # highcharts - 1

    elem = WebDriverWait(browser, 30).until(EC.element_to_be_clickable(
        (By.XPATH, '//*[@id="InputSectionBind"]/div/ul/li[1]/div[5]/a[1]')))
    time.sleep(5)
    elem.click()
    # browser.find_element_by_xpath('//*[@id="InputSectionBind"]/div/ul/li[1]/div[5]/a[1]').click()
    # time.sleep(5)
    WebDriverWait(browser, 10).until(EC.element_to_be_clickable(
        (By.XPATH, '//*[@id="EmployeeCreationFileUploadSection"]/label/i'))).click()
    # browser.find_element_by_xpath('//*[@id="EmployeeCreationFileUploadSection"]/label/i').click()
    time.sleep(40)


########################uncomment to upload data############################33
# autoit.control_focus("Open", "Edit1")
# fname = os.path.join(os.getcwd(),"Excel Files","Zing\Zing(13-06-2019).xlsx")
# autoit.control_set_text("Open", "Edit1", fname)
# autoit.control_click("Open", "Button1")
####################################################33
# """

def save_photo(link, save_path):
    fid = str(link).split("id=")[1].strip()
    gdd.download_file_from_google_drive(file_id=fid,
                                        dest_path=save_path)


def On_Contract(record, path):
    doj = record[58]
    doj = doj[3:5] + '/' + doj[0:2] + '/' + doj[-4:]

    doc = record[62]
    doc = doc[3:5] + '/' + doc[0:2] + '/' + doc[-4:]

    d1 = datetime.datetime.strptime(doc, "%d/%m/%Y")
    d2 = datetime.datetime.strptime(doj, "%d/%m/%Y")
    months = math.ceil((abs((d2 - d1).days) - 1) / 30)

    document = Document()

    p = document.add_paragraph(F"""Date: {doj}

	{record[1]} {record[2]} {record[3]} {record[4]} ,
	{record[12]},
	City: {record[13]}
	State: {record[14]}
	Mobile No.: {record[16]}
	Personal email ID: {record[35]}

	Dear {record[2]},

	We are pleased to offer you fixed term employment contract at Mahindra Integrated Business Solutions Pvt. Ltd. (“MIBS”) as per the terms stated herewith:""")
    DEPUTATION = document.add_paragraph('')
    DEPUTATION.add_run('DEPUTATION:').bold = True
    document.add_paragraph(
        F"""You are deputed to {record[88]} (“Client”) office under this contract. The terms of employment is exclusively with MIBS, the employee shall never be deemed to be the employee of the Client, where you have been deputed under this Contract. 

	You will with effect from {doj} be deputed at the Client’s office at any of their locations.""")

    TENURE = document.add_paragraph('')
    TENURE.add_run("TENURE:").bold = True
    document.add_paragraph(
        F"""The term of your Contract shall be valid for a period of {months} Months from {doj} to {doc}.

	In the event of Contract between Client and MIBS terminates for the project/work/deputation for which you are being employed, this Contract of employment shall be terminated.""")

    EXTENSION = document.add_paragraph('')
    EXTENSION.add_run('EXTENSION:').bold = True
    document.add_paragraph(
        F"""Unless otherwise notified to you in writing, this contract of employment will be valid till {record[62]}. This Contract may be considered for extension depending on the Client’s and MIBS requirement. The extension of Contract would be considered on fresh terms as agreed between you and MIBS through a separate mutually executed contract of employment. MIBS shall inform you in writing of the extension requirements.""")

    LOCATION = document.add_paragraph('')
    LOCATION.add_run('LOCATION:').bold = True
    document.add_paragraph(F"""You are required to work at the Client’s following office: {record[70]}.""")

    POSITION = document.add_paragraph('')
    POSITION.add_run('POSITION:').bold = True
    document.add_paragraph(F"""You are being appointed as {record[67]}.""")

    REMUNERATION = document.add_paragraph('')
    REMUNERATION.add_run('REMUNERATION:').bold = True
    document.add_paragraph(
        F"""The details of your salary break up with components are mentioned in Annexure 1 to this document""")

    WORKING_HOURS = document.add_paragraph('')
    WORKING_HOURS.add_run('WORKING HOURS:').bold = True
    document.add_paragraph(
        F"""You will follow the working hours of the Client where you will be deputed. You may have to work in shifts basis Client’s requirements. Your attendance will be maintained by the Reporting Supervisor of the Client which needs to be mandatorily sent to MIBS on or before the cutoff date mutually agreed between Client and MIBS for salary processing.""")

    TERMINATION_AND_SUSPENSION = document.add_paragraph('')
    TERMINATION_AND_SUSPENSION.add_run('TERMINATION AND SUSPENSION:').bold = True
    document.add_paragraph(
        F"""At the time of termination of employment either due to termination by your or MIBS or upon lapse of the term of employment, if there are any dues owing from you to MIBS, the same may be adjusted against any monies due to you by MIBS on account of salary including bonus and any other payment due to you under the terms of employment.

	During the tenure of employment any deviation or misconduct in any form noticed by MIBS or if there are any breach of internal policies pertaining to MIBS or the Client or any policy and regulation that was mutually agreed to be complied with, MIBS or Principal Employer has the right to suspend your services until you are notified to resume services in writing by MIBS. MIBS has the right to withhold full or part of your salary during this suspension period.""")

    NOTICE_PERIOD = document.add_paragraph('')
    NOTICE_PERIOD.add_run('NOTICE PERIOD:').bold = True
    document.add_paragraph(
        F"""The Company will be at liberty to terminate your services with 30 days’ notice or by paying you 30 days’ salary, including allowances, in lieu of notice.  In the event the Company decides to pay you 30 days’ salary in lieu of notice, the Company will be at liberty to call upon you not to take up any alternate employment for the period of 30 days. The Company will also be at liberty to call upon you not to report for work, though you would be on the rolls of the Company for the said period and you would be paid your salary as per your contract, as if you were on duty. In the event you choose to resign from the services of the Company, you will be required to serve for the period of notice of 30 days’.  The Company, however, will be at liberty to call upon you not to report for work or even take up any alternate employment during this period, which will be at the sole discretion of the Company.  The Company will also be at liberty to pay you 30 days’ notice wages in lieu of notice. However, it will be impermissible for you to waive the shortfall in the notice period by buying the said shortfall period in lieu thereof except with written permission.""")

    INDEMNITY = document.add_paragraph('')
    INDEMNITY.add_run('INDEMNITY:').bold = True
    document.add_paragraph(
        F"""You shall be responsible for protecting any property of the Client entrusted to you during the discharge of your duties and you shall indemnify the Client if there is any loss to the said property.""")

    NON_DISCLOSURE = document.add_paragraph('')
    NON_DISCLOSURE.add_run('NON DISCLOSURE:').bold = True
    document.add_paragraph(
        F"""You shall at all times endeavor to protect Proprietary and Confidential information, electronic information, electronic form, electronic record, information within your purview, systems & processes from unauthorized disclosure, alteration or destruction.""")

    CODE_CONDUCT = document.add_paragraph('')
    CODE_CONDUCT.add_run('CODE OF CONDUCT:').bold = True
    document.add_paragraph(
        F"""You shall not engage in any act subversive of discipline in the course of your duties for the Client either within the Client’s organization or outside it, and if are at any time found indulging in such acts, MIBS reserves the right to initiate disciplinary action against you as it deems fit.""")

    LEAVES_HOLIDAYS = document.add_paragraph('')
    LEAVES_HOLIDAYS.add_run('LEAVES / HOLIDAYS:').bold = True
    document.add_paragraph(
        F"""You shall be entitled to avail of the applicable holidays and weekly offs as may be applicable at your working location.

	Leaves entitlement will be as per the leave policy provided by MIBS.""")

    ADDRESS_COMMUNICATION = document.add_paragraph('')
    ADDRESS_COMMUNICATION.add_run('ADDRESS FOR COMMUNICATION:').bold = True
    document.add_paragraph(
        F"""The address for communication for the purpose of service of notice and for other official communication to MIBS shall be at the registered office of MIBS. The address for communication for service of notice and for other official communication to you shall be at the address stated in this document as your permanent address. In the event of any change in your address, you shall inform the same in writing to MIBS and that shall be the last address furnished by you, shall be deemed to be sufficient for communication and shall be deemed to be effective on you.""")

    BACKGROUND_VERIFICATION = document.add_paragraph('')
    BACKGROUND_VERIFICATION.add_run('BACKGROUND VERIFICATION:').bold = True
    document.add_paragraph(
        F"""MIBS reserves the right to have your background verified directly or through an outside agency. If on such verification it is found that you have furnished incorrect information or concealed material information, your services are liable to be terminated with immediate effect.""")

    ABSENTEEISM = document.add_paragraph('')
    ABSENTEEISM.add_run('ABSENTEEISM:').bold = True
    document.add_paragraph(
        F"""You should be regular and punctual in attendance. If you remain absent for 3 consecutive working days or more without sanction of leave or prior permission of if you overstayed your sanctioned leave beyond 3 working days or more, it shall be deemed that you have voluntarily abandoned your employment with MIBS and your services are liable to be terminated accordingly.""")

    RULES_REGULATIONS = document.add_paragraph('')
    RULES_REGULATIONS.add_run('RULES AND REGULATIONS:').bold = True
    document.add_paragraph(
        F"""You shall be bound by the rules and regulations framed by MIBS from time to time in relation to conduct, discipline and other service conditions which will be deemed as Rules and 
	Regulations and shall form part and parcel of this letter of appointment.""")

    OTHER_CONTRACT = document.add_paragraph('')
    OTHER_CONTRACT.add_run('OTHER TERMS OF CONTRACT:').bold = True
    document.add_paragraph(
        F"""In addition to the terms of appointment mentioned above, you are also governed by the standard employment rules of MIBS. The combined rules and procedures as contained in this letter will constitute the standard employment rules and you are required to read both of them together.

	During the Contract Period, you will not engage yourself directly or indirectly in other business or work part-time or accept any other form of employment or contract assignments except to the extent as may be permitted by us in writing.""")

    JURISDICTION = document.add_paragraph('')
    JURISDICTION.add_run('JURISDICTION:').bold = True
    document.add_paragraph(
        F"""Notwithstanding the place of work or place of residence of employee or the place where this document has been signed or executed, this Contract shall only be subject to the jurisdiction of the Courts of Mumbai.""")

    DEEMED_EMPLOYMENT = document.add_paragraph('')
    DEEMED_EMPLOYMENT.add_run('DEEMED CANCELLATION OF EMPLOYMENT:').bold = True
    document.add_paragraph(
        F"""The offer for employment stands cancelled and revoked if you do not report to duty within 3 days from the date of joining and your act will be construed as deemed and implied rejection of the offer for employment from your side, hence no obligation would arise on the part of MIBS in lieu of this offer for employment.""")

    DOCUMENTS = document.add_paragraph('')
    DOCUMENTS.add_run('DOCUMENTS:').bold = True
    document.add_paragraph(F"""You are required to submit the following documents on the date of joining:""")

    document.add_paragraph('Educational certificate', style='List Number')
    document.add_paragraph('Relieving/Experience letters from previous employers', style='List Number')
    document.add_paragraph('Last 3 months salary slip', style='List Number')
    document.add_paragraph('Identity proof with photograph', style='List Number')
    document.add_paragraph('Address proof – permanent and communication address', style='List Number')
    document.add_paragraph('3 passport size photographs', style='List Number')
    document.add_paragraph('PAN card copy', style='List Number')
    document.add_paragraph('UAN card copy', style='List Number')
    document.add_paragraph('Aadhar card copy', style='List Number')

    document.add_paragraph("""Wishing you the best in your assignment. As a token of understanding and accepting the terms of employment, you are requested to sign the duplicate copy of this letter and return to us within a day.

	Warm Regards,


	For Mahindra Integrated Business Solutions Pvt. Ltd.""")
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')
    sign = document.add_paragraph('')
    sign.add_run("""Riten Chakrabarty
	CFO""").bold = True
    document.add_paragraph("""I have read and understood the above mentioned terms and conditions of the Contract. I voluntarily accept the same. I shall abide by the terms and conditions stated by MIBS and any amendments from time to time. 

	All the above mentioned terms and conditions will come into force from your date of joining and in conditionsase no acceptance is received before the first salary it would be deemed as acknowledged and accepted by you on receipt of your first salary.""")
    document.add_paragraph('')
    line = document.add_paragraph('')
    line.add_run(
        """_________________________________________________________________________________________________________""").bold = True
    document.add_paragraph('')
    name = document.add_paragraph('')
    name.add_run('Name:').bold = True

    Signature = document.add_paragraph('')
    Signature.add_run('Signature:').bold = True
    document.add_paragraph('')

    place = document.add_paragraph('')
    place.add_run('Place:').bold = True

    date = document.add_paragraph('')
    date.add_run('Date:').bold = True

    document.save(path)


def On_rolls(record, path):
    doj = record[58]
    doj = doj[3:5] + '/' + doj[0:2] + '/' + doj[-4:]

    doc = record[62]
    doc = doc[3:5] + '/' + doc[0:2] + '/' + doc[-4:]

    dob = record[28]
    dob = dob[3:5] + '/' + dob[0:2] + '/' + dob[-4:]

    d1 = datetime.datetime.strptime(doc, "%d/%m/%Y")
    d2 = datetime.datetime.strptime(doj, "%d/%m/%Y")
    months = math.ceil((abs((d2 - d1).days) - 1) / 30)

    document = Document()

    p = document.add_paragraph(F"""Date: {doj}

	{record[1]} {record[2]} {record[3]} {record[4]} ,
	{record[12]},
	City: {record[13]}
	State: {record[14]}
	Mobile No.: {record[16]}
	Personal email ID: {record[35]}

	Dear {record[2]},

	We are pleased to offer you fixed term employment contract at Mahindra Integrated Business Solutions Pvt. Ltd. (“MIBS”) as per the terms stated herewith:""")
    DEPUTATION = document.add_paragraph('')
    DEPUTATION.add_run('DEPUTATION:').bold = True
    document.add_paragraph(
        F"""You are deputed to {record[88]} (“Client”) office under this contract. The terms of employment is exclusively with MIBS, the employee shall never be deemed to be the employee of the Client, where you have been deputed under this Contract. 

	You will with effect from {doj} be deputed at the Client’s office at any of their locations.""")

    LOCATION = document.add_paragraph('')
    LOCATION.add_run('LOCATION:').bold = True
    document.add_paragraph(F"""You are required to work at the Client’s following office: {record[70]}.""")

    POSITION = document.add_paragraph('')
    POSITION.add_run('POSITION:').bold = True
    document.add_paragraph(F"""You are being appointed as {record[67]}.""")

    REMUNERATION = document.add_paragraph('')
    REMUNERATION.add_run('REMUNERATION:').bold = True
    document.add_paragraph(
        F"""The details of your salary break up with components are mentioned in Annexure 1 to this document""")

    WORKING_HOURS = document.add_paragraph('')
    WORKING_HOURS.add_run('WORKING HOURS:').bold = True
    document.add_paragraph(
        F"""You will follow the working hours of the Client where you will be deputed. You may have to work in shifts basis Client’s requirements. Your attendance will be maintained by the Reporting Supervisor of the Client which needs to be mandatorily sent to MIBS on or before the cutoff date mutually agreed between Client and MIBS for salary processing.""")

    TERMINATION_AND_SUSPENSION = document.add_paragraph('')
    TERMINATION_AND_SUSPENSION.add_run('TERMINATION AND SUSPENSION:').bold = True
    document.add_paragraph(
        F"""At the time of termination of employment either due to termination by your or MIBS or upon lapse of the term of employment, if there are any dues owing from you to MIBS, the same may be adjusted against any monies due to you by MIBS on account of salary including bonus and any other payment due to you under the terms of employment.

	During the tenure of employment any deviation or misconduct in any form noticed by MIBS or if there are any breach of internal policies pertaining to MIBS or the Client or any policy and regulation that was mutually agreed to be complied with, MIBS or Principal Employer has the right to suspend your services until you are notified to resume services in writing by MIBS. MIBS has the right to withhold full or part of your salary during this suspension period.""")

    NOTICE_PERIOD = document.add_paragraph('')
    NOTICE_PERIOD.add_run('NOTICE PERIOD:').bold = True
    document.add_paragraph(
        F"""The Company will be at liberty to terminate your services with 30 days’ notice or by paying you 30 days’ salary, including allowances, in lieu of notice.  In the event the Company decides to pay you 30 days’ salary in lieu of notice, the Company will be at liberty to call upon you not to take up any alternate employment for the period of 30 days. The Company will also be at liberty to call upon you not to report for work, though you would be on the rolls of the Company for the said period and you would be paid your salary as per your contract, as if you were on duty. In the event you choose to resign from the services of the Company, you will be required to serve for the period of notice of 30 days’.  The Company, however, will be at liberty to call upon you not to report for work or even take up any alternate employment during this period, which will be at the sole discretion of the Company.  The Company will also be at liberty to pay you 30 days’ notice wages in lieu of notice. However, it will be impermissible for you to waive the shortfall in the notice period by buying the said shortfall period in lieu thereof except with written permission.""")

    INDEMNITY = document.add_paragraph('')
    INDEMNITY.add_run('INDEMNITY:').bold = True
    document.add_paragraph(
        F"""You shall be responsible for protecting any property of the Client entrusted to you during the discharge of your duties and you shall indemnify the Client if there is any loss to the said property.""")

    NON_DISCLOSURE = document.add_paragraph('')
    NON_DISCLOSURE.add_run('NON DISCLOSURE:').bold = True
    document.add_paragraph(
        F"""You shall at all times endeavor to protect Proprietary and Confidential information, electronic information, electronic form, electronic record, information within your purview, systems & processes from unauthorized disclosure, alteration or destruction.""")

    CODE_CONDUCT = document.add_paragraph('')
    CODE_CONDUCT.add_run('CODE OF CONDUCT:').bold = True
    document.add_paragraph(
        F"""You shall not engage in any act subversive of discipline in the course of your duties for the Client either within the Client’s organization or outside it, and if are at any time found indulging in such acts, MIBS reserves the right to initiate disciplinary action against you as it deems fit.""")

    LEAVES_HOLIDAYS = document.add_paragraph('')
    LEAVES_HOLIDAYS.add_run('LEAVES / HOLIDAYS:').bold = True
    document.add_paragraph(
        F"""You shall be entitled to avail of the applicable holidays and weekly offs as may be applicable at your working location.

	Leaves entitlement will be as per the leave policy provided by MIBS.""")

    HOLIDAYS = document.add_paragraph('')
    HOLIDAYS.add_run('HOLIDAYS:').bold = True
    document.add_paragraph(
        F"""You will be entitled to paid holidays in a year as notified by MIBS from time to time.""")

    ADDRESS_COMMUNICATION = document.add_paragraph('')
    ADDRESS_COMMUNICATION.add_run('ADDRESS FOR COMMUNICATION:').bold = True
    document.add_paragraph(
        F"""The address for communication for the purpose of service of notice and for other official communication to MIBS shall be at the registered office of MIBS. The address for communication for service of notice and for other official communication to you shall be at the address stated in this document as your permanent address. In the event of any change in your address, you shall inform the same in writing to MIBS and that shall be the last address furnished by you, shall be deemed to be sufficient for communication and shall be deemed to be effective on you.""")

    BACKGROUND_VERIFICATION = document.add_paragraph('')
    BACKGROUND_VERIFICATION.add_run('BACKGROUND VERIFICATION:').bold = True
    document.add_paragraph(
        F"""MIBS reserves the right to have your background verified directly or through an outside agency. If on such verification it is found that you have furnished incorrect information or concealed material information, your services are liable to be terminated with immediate effect.""")

    RETIREMENT_AGE = document.add_paragraph('')
    RETIREMENT_AGE.add_run('RETIREMENT AGE:').bold = True
    document.add_paragraph(
        F"""The age of retirement will be sixty years. (On the strength of the bio-data submitted to you, we have recorded your date of birth as {dob}.""")

    ABSENTEEISM = document.add_paragraph('')
    ABSENTEEISM.add_run('ABSENTEEISM:').bold = True
    document.add_paragraph(
        F"""You should be regular and punctual in attendance. If you remain absent for 3 consecutive working days or more without sanction of leave or prior permission of if you overstayed your sanctioned leave beyond 3 working days or more, it shall be deemed that you have voluntarily abandoned your employment with MIBS and your services are liable to be terminated accordingly.""")

    RULES_REGULATIONS = document.add_paragraph('')
    RULES_REGULATIONS.add_run('RULES AND REGULATIONS:').bold = True
    document.add_paragraph(
        F"""You shall be bound by the rules and regulations framed by MIBS from time to time in relation to conduct, discipline and other service conditions which will be deemed as Rules and 
	Regulations and shall form part and parcel of this letter of appointment.""")

    OTHER_CONTRACT = document.add_paragraph('')
    OTHER_CONTRACT.add_run('OTHER TERMS OF CONTRACT:').bold = True
    document.add_paragraph(
        F"""In addition to the terms of appointment mentioned above, you are also governed by the standard employment rules of MIBS. The combined rules and procedures as contained in this letter will constitute the standard employment rules and you are required to read both of them together.

	During the Contract Period, you will not engage yourself directly or indirectly in other business or work part-time or accept any other form of employment or contract assignments except to the extent as may be permitted by us in writing.""")

    JURISDICTION = document.add_paragraph('')
    JURISDICTION.add_run('JURISDICTION:').bold = True
    document.add_paragraph(
        F"""Notwithstanding the place of work or place of residence of employee or the place where this document has been signed or executed, this Contract shall only be subject to the jurisdiction of the Courts of Mumbai.""")

    DEEMED_EMPLOYMENT = document.add_paragraph('')
    DEEMED_EMPLOYMENT.add_run('DEEMED CANCELLATION OF EMPLOYMENT:').bold = True
    document.add_paragraph(
        F"""The offer for employment stands cancelled and revoked if you do not report to duty within 3 days from the date of joining and your act will be construed as deemed and implied rejection of the offer for employment from your side, hence no obligation would arise on the part of MIBS in lieu of this offer for employment.""")

    DOCUMENTS = document.add_paragraph('')
    DOCUMENTS.add_run('DOCUMENTS:').bold = True
    document.add_paragraph(F"""You are required to submit the following documents on the date of joining:""")

    document.add_paragraph('Educational certificate', style='List Number')
    document.add_paragraph('Relieving/Experience letters from previous employers', style='List Number')
    document.add_paragraph('Last 3 months salary slip', style='List Number')
    document.add_paragraph('Identity proof with photograph', style='List Number')
    document.add_paragraph('Address proof – permanent and communication address', style='List Number')
    document.add_paragraph('3 passport size photographs', style='List Number')
    document.add_paragraph('PAN card copy', style='List Number')
    document.add_paragraph('UAN card copy', style='List Number')
    document.add_paragraph('Aadhar card copy', style='List Number')

    document.add_paragraph("""Wishing you the best in your assignment. As a token of understanding and accepting the terms of employment, you are requested to sign the duplicate copy of this letter and return to us within a day.

	Warm Regards,


	For Mahindra Integrated Business Solutions Pvt. Ltd.""")
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')
    sign = document.add_paragraph('')
    sign.add_run("""Riten Chakrabarty
	CFO""").bold = True
    document.add_paragraph("""I have read and understood the above mentioned terms and conditions of the Contract. I voluntarily accept the same. I shall abide by the terms and conditions stated by MIBS and any amendments from time to time. 

	All the above mentioned terms and conditions will come into force from your date of joining and in conditionsase no acceptance is received before the first salary it would be deemed as acknowledged and accepted by you on receipt of your first salary.""")
    document.add_paragraph('')
    line = document.add_paragraph('')
    line.add_run(
        """_________________________________________________________________________________________________________""").bold = True
    document.add_paragraph('')
    name = document.add_paragraph('')
    name.add_run('Name:').bold = True

    Signature = document.add_paragraph('')
    Signature.add_run('Signature:').bold = True
    document.add_paragraph('')

    place = document.add_paragraph('')
    place.add_run('Place:').bold = True

    date = document.add_paragraph('')
    date.add_run('Date:').bold = True

    document.save(path)


def L_Band(record, path):
    doj = record[58]
    doj = doj[3:5] + '/' + doj[0:2] + '/' + doj[-4:]

    dob = record[28]
    dob = dob[3:5] + '/' + dob[0:2] + '/' + dob[-4:]

    document = Document()

    p = document.add_paragraph(F"""





	{doj}

	Employee Code: {record[57]}

	{record[1]} {record[2]} {record[3]} {record[4]}
	{record[12]},
	{record[13]},{record[15]}

	Dear {record[2]},

	We have pleasure of offering you appointment in our company, as {record[67]} – {record[84]} in Band {record[66]}.

	This offer of appointment is subject to you satisfying the following:""")
    document.add_paragraph(
        F"""Your written undertaking to join, not later than {doj}. You have been deputed to work at {record[70]}.""",
        style='List Bullet')
    document.add_paragraph(
        F"""Submission of all your necessary certificates and documents in respect of educational and professional qualifications, proof of age and previous employment, as per the requirements of the Company.""",
        style='List Bullet')

    document.add_paragraph(F"""This appointment will commence from the date on which you actually join the Company. 

	The terms and conditions of employment shall be as under:""")

    document.add_paragraph("""You will receive emoluments / allowances as per the attached Annexure.""",
                           style='List Number')
    a = document.add_paragraph("""Your employment will be on the rolls of """, style='List Number')
    a.add_run("'Mahindra Integrated Business Solutions Pvt. Ltd.'").bold = True

    document.add_paragraph(
        """Gratuity, Provident Fund, and Medical Benefit and family, as per the rules of the Company, Medical Benefit will be applicable from the date of confirmation in service.""",
        style='List Number')

    document.add_paragraph(
        F"""The age of retirement will be sixty years. (On the strength of the bio-data submitted by you, we have recorded your date of birth as {dob})""",
        style='List Number')
    document.add_paragraph(
        """With effect from the date of your employment, you are required to become a member Provident Fund.""",
        style='List Number')
    document.add_paragraph(
        """You are liable to be placed for service at our discretion at any of the Company's establishments/departments/divisions anywhere in India as also the Associate Companies and subsidiary Companies of Mahindra & Mahindra Ltd.""",
        style='List Number')
    document.add_paragraph(
        """You will be on probation for a period of 06 months. This probationary period could be curtailed or extended at the discretion of the Company. However, completion of 06 months of probation does not entitle you or result in automatic confirmation of your employment, unless the Company confirms your employment in writing. During this period, your employment may be terminated forthwith without notice and without assigning any reasons.""",
        style='List Number')
    document.add_paragraph(
        """You will be entitled to Exigency Leave as per the rules. Upon satisfactory completion of the period of your probation, you will be confirmed in our service and on confirmation:""",
        style='List Number')
    document.add_paragraph("""You will be entitled to Privilege Leave, as per the rules thereof;""",
                           style='List Bullet')
    document.add_paragraph(
        """The Company will be at liberty to terminate your services with 90 days notice or by paying you 90 days salary, including allowances, in lieu of notice. In the event the Company decides to pay you 90 days salary in lieu of notice, the Company will be at liberty to call upon you not to take up any alternate employment for the period of 90 days. The Company will also be at liberty to call upon you not to report for work, though you would be on the rolls of the Company for the said period and you would be paid your salary as per your contract, as if you were on duty. In the event you choose to resign from the services of the Company, you will be required to serve for the period of notice of 90 days. The Company, however, will be at liberty to call upon you not to report for work or even take up any alternate employment during this period, which will be at the sole discretion of the Company. The Company will also be at liberty to pay you 90 days notice wages in lieu of notice. However, it will be impermissible for you to waive the shortfall in the notice period by buying the said shortfall period in lieu thereof except with written permission.""",
        style='List Bullet')
    document.add_paragraph(
        """During the probationary period, the notice period to be served by either party shall be 15 days.""",
        style='List Bullet')
    document.add_paragraph(
        """Further, you shall not be entitled to adjust your notice period against privilege leave, if any, standing to you credit.""",
        style='List Bullet')
    document.add_paragraph(
        """So long as you are in the employment of the Company, you will, at all times, observe secrecy and confidentiality and will not divulge, disclose or make known to any unauthorized person within or outside the Company, nor will you unauthorized use any knowledge or information in respect of manufacturing, technical trade or business data (including manufacturing processes, technical know-how, customer information, business plans and like matters) which are necessarily confidential and have come to your knowledge and possession.  You will also not remove any such information in any form whatsoever from the Company premises, nor copy or transmit the same unauthorized nor will you grant permission to assist, permit entry to, or in any manner co-operate with any unauthorized person for the purposes of accessing, obtaining, copying, transmitting or removing the above. Even after the cessation of your employment with the Company, you will not use, divulge, disclose or remove in any manner whatsoever confidential information of the type described above of which you were in possession whilst in service to the detriment of the Company.   You will also observe all the confidentiality measures which are in existence, or which may be enforced from time to time, as well as directions as to confidentiality marked on any communication, document, computer floppy etc. You shall indemnify and hold Company harmless and indemnified against any damage or loss caused to the Company on account of breach of confidentiality on your part. These confidentiality provisions shall survive the separation of your employment with the Company, either by way of retirement or termination or otherwise.""",
        style='List Number')
    document.add_paragraph(
        """In addition to your fulfilling the requirements of secrecy and confidentiality, as specified herein, also during your employment with the Company, you shall not engage in any vocation, training, employment, consultancy, business, transaction, or any other activity, which is in conflict with the interests of the Company, in any capacity whatsoever either on your own or in association with any other individual/firm/institute/body corporate, etc., whether for any consideration or not.""",
        style='List Number')
    document.add_paragraph(
        """You will devote your full attention exclusively to the duties entrusted to you from time to time by the Company and while in service of this Company you will not work for any person or Company in any capacity either for any consideration or otherwise, nor do any private business without obtaining prior permission of the Company in writing.""",
        style='List Number')
    document.add_paragraph(
        """You will assign to the Company your entire right, title and interest in any Intellectual Property Rights (IPRs for short, which term would include patents, trade-marks, copyrights, designs, whether registered or not, and all improvements thereto) that you may make, solely or jointly with others, in the course of your employment with the Company relating to any or all systems, services and products manufactured or marketed or leased or developed.  You  will perform all necessary  acts and  execute such documents in such format as may be required by the Company,  without  expense  to you, which in the  judgment  of  the  Company  or its Attorneys may be necessary or desirable to  secure  to the Company full right title and interest in the IPRs.""",
        style='List Number')
    document.add_paragraph(
        """The Company shall at all times have the right to access and monitor all e-mails created, sent / received or stored by you using Company facility and on Company’s system at any time without giving you any prior notification. All such data and information shall be the property of the Company at all times.""",
        style='List Number')
    document.add_paragraph(
        """You shall endeavor to uphold the good image of the Company and shall not by your conduct adversely affect the reputation of the Company and bring disrepute to the Company, in any manner whatsoever.""",
        style='List Number')
    document.add_paragraph(
        """You shall, on ceasing to be the employee of the Company, forthwith return all Company properties, movable and immovable, including, without limitation, all Company information, files, reports, memoranda, software, credit cards, door and file keys, computer access codes and such other property which you received or in possession or prepared in connection with your employment with the Company""",
        style='List Number')
    document.add_paragraph(
        """Any joining expenses reimbursed by the Company will be recovered in the event you leave the organization within one year of joining.""",
        style='List Number')
    document.add_paragraph(
        """You will be subject to all rules, regulations and policies of the Company, which may be in force from time to time.""",
        style='List Number')
    document.add_paragraph(
        """Any joining expenses / relocation expenses / Notice Pay buyout reimbursed or any other payments made to you while joining by the Company will be recovered in the event you leave the organization within one year of joining.""",
        style='List Number')
    document.add_paragraph("""Please return the duplicate of this letter, duly signed, in token of your acceptance of the above mentioned terms and conditions of the employment, having read the attached Code of Conduct for Senior Management & Employees and on joining you will abide by its prescriptive principles.


	We wish you a long and fruitful career with us. 

	With Regards,

	Yours Sincerely,""")
    sign = document.add_paragraph('')
    sign.add_run("""For MAHINDRA INTEGRATED BUSINESS SOLUTIONS PRIVATE LIMITED





	Riten Chakrabarty
	Chief Finance Officer.""").bold = True

    document.save(path)


def MB_and_Above(record, path):
    doj = record[58]
    doj = doj[3:5] + '/' + doj[0:2] + '/' + doj[-4:]

    dob = record[28]
    dob = dob[3:5] + '/' + dob[0:2] + '/' + dob[-4:]

    document = Document()

    p = document.add_paragraph(F"""





	{doj}

	Employee Code: {record[57]}

	{record[1]} {record[2]} {record[3]} {record[4]}
	{record[12]},
	{record[13]},{record[15]}

	Dear {record[2]},

	We have pleasure of offering you appointment in our company, as {record[67]} – {record[84]} in Band {record[66]}.

	This offer of appointment is subject to you satisfying the following:""")
    document.add_paragraph(
        F"""Your written undertaking to join, not later than {doj}. You have been deputed to work at {record[70]}.""",
        style='List Bullet')
    document.add_paragraph(
        F"""Submission of all your necessary certificates and documents in respect of educational and professional qualifications, proof of age and previous employment, as per the requirements of the Company.""",
        style='List Bullet')

    document.add_paragraph(F"""This appointment will commence from the date on which you actually join the Company. 

	The terms and conditions of employment shall be as under:""")

    document.add_paragraph("""You will receive emoluments / allowances as per the attached Annexure.""",
                           style='List Number')
    a = document.add_paragraph("""Your employment will be on the rolls of """, style='List Number')
    a.add_run("'Mahindra Integrated Business Solutions Pvt. Ltd.'").bold = True

    document.add_paragraph(
        """Gratuity, Provident Fund, and Medical Benefit and family, as per the rules of the Company, Medical Benefit will be applicable from the date of confirmation in service.""",
        style='List Number')

    document.add_paragraph(
        F"""The age of retirement will be sixty years. (On the strength of the bio-data submitted by you, we have recorded your date of birth as {dob})""",
        style='List Number')
    document.add_paragraph(
        """With effect from the date of your employment, you are required to become a member Provident Fund.""",
        style='List Number')
    document.add_paragraph(
        """You are liable to be placed for service at our discretion at any of the Company's establishments/departments/divisions anywhere in India as also the Associate Companies and subsidiary Companies of Mahindra & Mahindra Ltd.""",
        style='List Number')
    document.add_paragraph(
        """You will be on probation for a period of 06 months. This probationary period could be curtailed or extended at the discretion of the Company. However, completion of 06 months of probation does not entitle you or result in automatic confirmation of your employment, unless the Company confirms your employment in writing. During this period, your employment may be terminated forthwith without notice and without assigning any reasons.""",
        style='List Number')
    document.add_paragraph(
        """You will be entitled to Exigency Leave as per the rules. Upon satisfactory completion of the period of your probation, you will be confirmed in our service and on confirmation:""",
        style='List Number')
    document.add_paragraph("""You will be entitled to Privilege Leave, as per the rules thereof;""",
                           style='List Bullet')
    document.add_paragraph(
        """The Company will be at liberty to terminate your services with 90 days notice or by paying you 90 days salary, including allowances, in lieu of notice. In the event the Company decides to pay you 90 days salary in lieu of notice, the Company will be at liberty to call upon you not to take up any alternate employment for the period of 90 days. The Company will also be at liberty to call upon you not to report for work, though you would be on the rolls of the Company for the said period and you would be paid your salary as per your contract, as if you were on duty. In the event you choose to resign from the services of the Company, you will be required to serve for the period of notice of 90 days. The Company, however, will be at liberty to call upon you not to report for work or even take up any alternate employment during this period, which will be at the sole discretion of the Company. The Company will also be at liberty to pay you 90 days notice wages in lieu of notice. However, it will be impermissible for you to waive the shortfall in the notice period by buying the said shortfall period in lieu thereof except with written permission.""",
        style='List Bullet')
    document.add_paragraph(
        """During the probationary period, the notice period to be served by either party shall be 15 days.""",
        style='List Bullet')
    document.add_paragraph(
        """Further, you shall not be entitled to adjust your notice period against privilege leave, if any, standing to you credit.""",
        style='List Bullet')
    document.add_paragraph(
        """So long as you are in the employment of the Company, you will, at all times, observe secrecy and confidentiality and will not divulge, disclose or make known to any unauthorized person within or outside the Company, nor will you unauthorized use any knowledge or information in respect of manufacturing, technical trade or business data (including manufacturing processes, technical know-how, customer information, business plans and like matters) which are necessarily confidential and have come to your knowledge and possession.  You will also not remove any such information in any form whatsoever from the Company premises, nor copy or transmit the same unauthorized nor will you grant permission to assist, permit entry to, or in any manner co-operate with any unauthorized person for the purposes of accessing, obtaining, copying, transmitting or removing the above. Even after the cessation of your employment with the Company, you will not use, divulge, disclose or remove in any manner whatsoever confidential information of the type described above of which you were in possession whilst in service to the detriment of the Company.   You will also observe all the confidentiality measures which are in existence, or which may be enforced from time to time, as well as directions as to confidentiality marked on any communication, document, computer floppy etc. You shall indemnify and hold Company harmless and indemnified against any damage or loss caused to the Company on account of breach of confidentiality on your part. These confidentiality provisions shall survive the separation of your employment with the Company, either by way of retirement or termination or otherwise.""",
        style='List Number')
    document.add_paragraph(
        """In addition to your fulfilling the requirements of secrecy and confidentiality, as specified herein, also during your employment with the Company, you shall not engage in any vocation, training, employment, consultancy, business, transaction, or any other activity, which is in conflict with the interests of the Company, in any capacity whatsoever either on your own or in association with any other individual/firm/institute/body corporate, etc., whether for any consideration or not.""",
        style='List Number')
    document.add_paragraph(
        """You will devote your full attention exclusively to the duties entrusted to you from time to time by the Company and while in service of this Company you will not work for any person or Company in any capacity either for any consideration or otherwise, nor do any private business without obtaining prior permission of the Company in writing.""",
        style='List Number')
    document.add_paragraph(
        """You will assign to the Company your entire right, title and interest in any Intellectual Property Rights (IPRs for short, which term would include patents, trade-marks, copyrights, designs, whether registered or not, and all improvements thereto) that you may make, solely or jointly with others, in the course of your employment with the Company relating to any or all systems, services and products manufactured or marketed or leased or developed.  You  will perform all necessary  acts and  execute such documents in such format as may be required by the Company,  without  expense  to you, which in the  judgment  of  the  Company  or its Attorneys may be necessary or desirable to  secure  to the Company full right title and interest in the IPRs.""",
        style='List Number')
    document.add_paragraph(
        """The Company shall at all times have the right to access and monitor all e-mails created, sent / received or stored by you using Company facility and on Company’s system at any time without giving you any prior notification. All such data and information shall be the property of the Company at all times.""",
        style='List Number')
    document.add_paragraph(
        """You shall endeavor to uphold the good image of the Company and shall not by your conduct adversely affect the reputation of the Company and bring disrepute to the Company, in any manner whatsoever.""",
        style='List Number')
    document.add_paragraph(
        """You shall, on ceasing to be the employee of the Company, forthwith return all Company properties, movable and immovable, including, without limitation, all Company information, files, reports, memoranda, software, credit cards, door and file keys, computer access codes and such other property which you received or in possession or prepared in connection with your employment with the Company""",
        style='List Number')
    document.add_paragraph(
        """Any joining expenses reimbursed by the Company will be recovered in the event you leave the organization within one year of joining.""",
        style='List Number')
    document.add_paragraph(
        """You will be subject to all rules, regulations and policies of the Company, which may be in force from time to time.""",
        style='List Number')
    document.add_paragraph(
        """Any joining expenses / relocation expenses / Notice Pay buyout reimbursed or any other payments made to you while joining by the Company will be recovered in the event you leave the organization within one year of joining.""",
        style='List Number')
    document.add_paragraph("""Please return the duplicate of this letter, duly signed, in token of your acceptance of the above mentioned terms and conditions of the employment, having read the attached Code of Conduct for Senior Management & Employees and on joining you will abide by its prescriptive principles.


	We wish you a long and fruitful career with us. 

	With Regards,

	Yours Sincerely,""")
    sign = document.add_paragraph('')
    sign.add_run("""For MAHINDRA INTEGRATED BUSINESS SOLUTIONS PRIVATE LIMITED





	Riten Chakrabarty
	Chief Finance Officer.""").bold = True

    document.save(path)


def MT(record, path):
    doj = record[58]
    doj = doj[3:5] + '/' + doj[0:2] + '/' + doj[-4:]

    dob = record[28]
    dob = dob[3:5] + '/' + dob[0:2] + '/' + dob[-4:]

    document = Document()

    p = document.add_paragraph(F"""





	{doj}

	Employee Code: {record[57]}

	{record[1]} {record[2]} {record[3]} {record[4]}
	{record[12]},
	{record[13]},{record[15]}

	Dear {record[2]},

	We have pleasure of offering you appointment in our company, as {record[67]} – {record[84]} in Band {record[66]}.

	This offer of appointment is subject to you satisfying the following:""")
    document.add_paragraph(
        F"""Your written undertaking to join, not later than {doj}. You have been deputed to work at {record[70]}.""",
        style='List Bullet')
    document.add_paragraph(
        F"""Submission of all your necessary certificates and documents in respect of educational and professional qualifications, proof of age and previous employment, as per the requirements of the Company.""",
        style='List Bullet')

    document.add_paragraph(F"""This appointment will commence from the date on which you actually join the Company. 

	The terms and conditions of employment shall be as under:""")

    document.add_paragraph("""You will receive emoluments / allowances as per the attached Annexure.""",
                           style='List Number')
    a = document.add_paragraph("""Your employment will be on the rolls of """, style='List Number')
    a.add_run("'Mahindra Integrated Business Solutions Pvt. Ltd.'").bold = True

    document.add_paragraph(
        """Gratuity, Provident Fund, and Medical Benefit and family, as per the rules of the Company, Medical Benefit will be applicable from the date of confirmation in service.""",
        style='List Number')

    document.add_paragraph(
        F"""The age of retirement will be sixty years. (On the strength of the bio-data submitted by you, we have recorded your date of birth as {dob})""",
        style='List Number')
    document.add_paragraph(
        """With effect from the date of your employment, you are required to become a member Provident Fund.""",
        style='List Number')
    document.add_paragraph(
        """You are liable to be placed for service at our discretion at any of the Company's establishments/departments/divisions anywhere in India as also the Associate Companies and subsidiary Companies of Mahindra & Mahindra Ltd.""",
        style='List Number')
    document.add_paragraph(
        """You will be on probation for a period of 06 months. This probationary period could be curtailed or extended at the discretion of the Company. However, completion of 06 months of probation does not entitle you or result in automatic confirmation of your employment, unless the Company confirms your employment in writing. During this period, your employment may be terminated forthwith without notice and without assigning any reasons.""",
        style='List Number')
    document.add_paragraph(
        """You will be entitled to Exigency Leave as per the rules. Upon satisfactory completion of the period of your probation, you will be confirmed in our service and on confirmation:""",
        style='List Number')
    document.add_paragraph("""You will be entitled to Privilege Leave, as per the rules thereof;""",
                           style='List Bullet')
    document.add_paragraph(
        """The Company will be at liberty to terminate your services with 90 days notice or by paying you 90 days salary, including allowances, in lieu of notice. In the event the Company decides to pay you 90 days salary in lieu of notice, the Company will be at liberty to call upon you not to take up any alternate employment for the period of 90 days. The Company will also be at liberty to call upon you not to report for work, though you would be on the rolls of the Company for the said period and you would be paid your salary as per your contract, as if you were on duty. In the event you choose to resign from the services of the Company, you will be required to serve for the period of notice of 90 days. The Company, however, will be at liberty to call upon you not to report for work or even take up any alternate employment during this period, which will be at the sole discretion of the Company. The Company will also be at liberty to pay you 90 days notice wages in lieu of notice. However, it will be impermissible for you to waive the shortfall in the notice period by buying the said shortfall period in lieu thereof except with written permission.""",
        style='List Bullet')
    document.add_paragraph(
        """During the probationary period, the notice period to be served by either party shall be 15 days.""",
        style='List Bullet')
    document.add_paragraph(
        """Further, you shall not be entitled to adjust your notice period against privilege leave, if any, standing to you credit.""",
        style='List Bullet')
    document.add_paragraph(
        """So long as you are in the employment of the Company, you will, at all times, observe secrecy and confidentiality and will not divulge, disclose or make known to any unauthorized person within or outside the Company, nor will you unauthorized use any knowledge or information in respect of manufacturing, technical trade or business data (including manufacturing processes, technical know-how, customer information, business plans and like matters) which are necessarily confidential and have come to your knowledge and possession.  You will also not remove any such information in any form whatsoever from the Company premises, nor copy or transmit the same unauthorized nor will you grant permission to assist, permit entry to, or in any manner co-operate with any unauthorized person for the purposes of accessing, obtaining, copying, transmitting or removing the above. Even after the cessation of your employment with the Company, you will not use, divulge, disclose or remove in any manner whatsoever confidential information of the type described above of which you were in possession whilst in service to the detriment of the Company.   You will also observe all the confidentiality measures which are in existence, or which may be enforced from time to time, as well as directions as to confidentiality marked on any communication, document, computer floppy etc. You shall indemnify and hold Company harmless and indemnified against any damage or loss caused to the Company on account of breach of confidentiality on your part. These confidentiality provisions shall survive the separation of your employment with the Company, either by way of retirement or termination or otherwise.""",
        style='List Number')
    document.add_paragraph(
        """In addition to your fulfilling the requirements of secrecy and confidentiality, as specified herein, also during your employment with the Company, you shall not engage in any vocation, training, employment, consultancy, business, transaction, or any other activity, which is in conflict with the interests of the Company, in any capacity whatsoever either on your own or in association with any other individual/firm/institute/body corporate, etc., whether for any consideration or not.""",
        style='List Number')
    document.add_paragraph(
        """You will devote your full attention exclusively to the duties entrusted to you from time to time by the Company and while in service of this Company you will not work for any person or Company in any capacity either for any consideration or otherwise, nor do any private business without obtaining prior permission of the Company in writing.""",
        style='List Number')
    document.add_paragraph(
        """You will assign to the Company your entire right, title and interest in any Intellectual Property Rights (IPRs for short, which term would include patents, trade-marks, copyrights, designs, whether registered or not, and all improvements thereto) that you may make, solely or jointly with others, in the course of your employment with the Company relating to any or all systems, services and products manufactured or marketed or leased or developed.  You  will perform all necessary  acts and  execute such documents in such format as may be required by the Company,  without  expense  to you, which in the  judgment  of  the  Company  or its Attorneys may be necessary or desirable to  secure  to the Company full right title and interest in the IPRs.""",
        style='List Number')
    document.add_paragraph(
        """The Company shall at all times have the right to access and monitor all e-mails created, sent / received or stored by you using Company facility and on Company’s system at any time without giving you any prior notification. All such data and information shall be the property of the Company at all times.""",
        style='List Number')
    document.add_paragraph(
        """You shall endeavor to uphold the good image of the Company and shall not by your conduct adversely affect the reputation of the Company and bring disrepute to the Company, in any manner whatsoever.""",
        style='List Number')
    document.add_paragraph(
        """You shall, on ceasing to be the employee of the Company, forthwith return all Company properties, movable and immovable, including, without limitation, all Company information, files, reports, memoranda, software, credit cards, door and file keys, computer access codes and such other property which you received or in possession or prepared in connection with your employment with the Company""",
        style='List Number')
    document.add_paragraph(
        """Any joining expenses reimbursed by the Company will be recovered in the event you leave the organization within one year of joining.""",
        style='List Number')
    document.add_paragraph(
        """You will be subject to all rules, regulations and policies of the Company, which may be in force from time to time.""",
        style='List Number')
    document.add_paragraph(
        """Any joining expenses / relocation expenses / Notice Pay buyout reimbursed or any other payments made to you while joining by the Company will be recovered in the event you leave the organization within one year of joining.""",
        style='List Number')
    document.add_paragraph("""Please return the duplicate of this letter, duly signed, in token of your acceptance of the above mentioned terms and conditions of the employment, having read the attached Code of Conduct for Senior Management & Employees and on joining you will abide by its prescriptive principles.


	We wish you a long and fruitful career with us. 

	With Regards,

	Yours Sincerely,""")
    sign = document.add_paragraph('')
    sign.add_run("""For MAHINDRA INTEGRATED BUSINESS SOLUTIONS PRIVATE LIMITED





	Riten Chakrabarty
	Chief Finance Officer.""").bold = True

    document.save(path)


def MS0_MS4(record, path):
    doj = record[58]
    doj = doj[3:5] + '/' + doj[0:2] + '/' + doj[-4:]

    dob = record[28]
    dob = dob[3:5] + '/' + dob[0:2] + '/' + dob[-4:]

    document = Document()

    p = document.add_paragraph(F"""





	{doj}

	Employee Code: {record[57]}

	{record[1]} {record[2]} {record[3]} {record[4]}
	{record[12]},
	{record[13]},{record[15]}

	Dear {record[2]},

	We have pleasure of offering you appointment in our company, as {record[67]} – {record[84]} in Band {record[66]}.

	This offer of appointment is subject to you satisfying the following:""")
    document.add_paragraph(
        F"""Your written undertaking to join, not later than {doj}. You have been deputed to work at {record[70]}.""",
        style='List Bullet')
    document.add_paragraph(
        F"""Submission of all your necessary certificates and documents in respect of educational and professional qualifications, proof of age and previous employment, as per the requirements of the Company.""",
        style='List Bullet')

    document.add_paragraph(F"""This appointment will commence from the date on which you actually join the Company. 

	The terms and conditions of employment shall be as under:""")

    document.add_paragraph("""You will receive emoluments / allowances as per the attached Annexure.""",
                           style='List Number')
    a = document.add_paragraph("""Your employment will be on the rolls of """, style='List Number')
    a.add_run("'Mahindra Integrated Business Solutions Pvt. Ltd.'").bold = True

    document.add_paragraph(
        """Gratuity, Provident Fund, and Medical Benefit and family, as per the rules of the Company, Medical Benefit will be applicable from the date of confirmation in service.""",
        style='List Number')

    document.add_paragraph(
        F"""The age of retirement will be sixty years. (On the strength of the bio-data submitted by you, we have recorded your date of birth as {dob})""",
        style='List Number')
    document.add_paragraph(
        """With effect from the date of your employment, you are required to become a member Provident Fund.""",
        style='List Number')
    document.add_paragraph(
        """You are liable to be placed for service at our discretion at any of the Company's establishments/departments/divisions anywhere in India as also the Associate Companies and subsidiary Companies of Mahindra & Mahindra Ltd.""",
        style='List Number')
    document.add_paragraph(
        """You will be on probation for a period of 06 months. This probationary period could be curtailed or extended at the discretion of the Company. However, completion of 06 months of probation does not entitle you or result in automatic confirmation of your employment, unless the Company confirms your employment in writing. During this period, your employment may be terminated forthwith without notice and without assigning any reasons.""",
        style='List Number')
    document.add_paragraph(
        """You will be entitled to Exigency Leave as per the rules. Upon satisfactory completion of the period of your probation, you will be confirmed in our service and on confirmation:""",
        style='List Number')
    document.add_paragraph("""You will be entitled to Privilege Leave, as per the rules thereof;""",
                           style='List Bullet')
    document.add_paragraph(
        """The Company will be at liberty to terminate your services with 90 days notice or by paying you 90 days salary, including allowances, in lieu of notice. In the event the Company decides to pay you 90 days salary in lieu of notice, the Company will be at liberty to call upon you not to take up any alternate employment for the period of 90 days. The Company will also be at liberty to call upon you not to report for work, though you would be on the rolls of the Company for the said period and you would be paid your salary as per your contract, as if you were on duty. In the event you choose to resign from the services of the Company, you will be required to serve for the period of notice of 90 days. The Company, however, will be at liberty to call upon you not to report for work or even take up any alternate employment during this period, which will be at the sole discretion of the Company. The Company will also be at liberty to pay you 90 days notice wages in lieu of notice. However, it will be impermissible for you to waive the shortfall in the notice period by buying the said shortfall period in lieu thereof except with written permission.""",
        style='List Bullet')
    document.add_paragraph(
        """During the probationary period, the notice period to be served by either party shall be 15 days.""",
        style='List Bullet')
    document.add_paragraph(
        """Further, you shall not be entitled to adjust your notice period against privilege leave, if any, standing to you credit.""",
        style='List Bullet')
    document.add_paragraph(
        """So long as you are in the employment of the Company, you will, at all times, observe secrecy and confidentiality and will not divulge, disclose or make known to any unauthorized person within or outside the Company, nor will you unauthorized use any knowledge or information in respect of manufacturing, technical trade or business data (including manufacturing processes, technical know-how, customer information, business plans and like matters) which are necessarily confidential and have come to your knowledge and possession.  You will also not remove any such information in any form whatsoever from the Company premises, nor copy or transmit the same unauthorized nor will you grant permission to assist, permit entry to, or in any manner co-operate with any unauthorized person for the purposes of accessing, obtaining, copying, transmitting or removing the above. Even after the cessation of your employment with the Company, you will not use, divulge, disclose or remove in any manner whatsoever confidential information of the type described above of which you were in possession whilst in service to the detriment of the Company.   You will also observe all the confidentiality measures which are in existence, or which may be enforced from time to time, as well as directions as to confidentiality marked on any communication, document, computer floppy etc. You shall indemnify and hold Company harmless and indemnified against any damage or loss caused to the Company on account of breach of confidentiality on your part. These confidentiality provisions shall survive the separation of your employment with the Company, either by way of retirement or termination or otherwise.""",
        style='List Number')
    document.add_paragraph(
        """In addition to your fulfilling the requirements of secrecy and confidentiality, as specified herein, also during your employment with the Company, you shall not engage in any vocation, training, employment, consultancy, business, transaction, or any other activity, which is in conflict with the interests of the Company, in any capacity whatsoever either on your own or in association with any other individual/firm/institute/body corporate, etc., whether for any consideration or not.""",
        style='List Number')
    document.add_paragraph(
        """You will devote your full attention exclusively to the duties entrusted to you from time to time by the Company and while in service of this Company you will not work for any person or Company in any capacity either for any consideration or otherwise, nor do any private business without obtaining prior permission of the Company in writing.""",
        style='List Number')
    document.add_paragraph(
        """You will assign to the Company your entire right, title and interest in any Intellectual Property Rights (IPRs for short, which term would include patents, trade-marks, copyrights, designs, whether registered or not, and all improvements thereto) that you may make, solely or jointly with others, in the course of your employment with the Company relating to any or all systems, services and products manufactured or marketed or leased or developed.  You  will perform all necessary  acts and  execute such documents in such format as may be required by the Company,  without  expense  to you, which in the  judgment  of  the  Company  or its Attorneys may be necessary or desirable to  secure  to the Company full right title and interest in the IPRs.""",
        style='List Number')
    document.add_paragraph(
        """The Company shall at all times have the right to access and monitor all e-mails created, sent / received or stored by you using Company facility and on Company’s system at any time without giving you any prior notification. All such data and information shall be the property of the Company at all times.""",
        style='List Number')
    document.add_paragraph(
        """You shall endeavor to uphold the good image of the Company and shall not by your conduct adversely affect the reputation of the Company and bring disrepute to the Company, in any manner whatsoever.""",
        style='List Number')
    document.add_paragraph(
        """You shall, on ceasing to be the employee of the Company, forthwith return all Company properties, movable and immovable, including, without limitation, all Company information, files, reports, memoranda, software, credit cards, door and file keys, computer access codes and such other property which you received or in possession or prepared in connection with your employment with the Company""",
        style='List Number')
    document.add_paragraph(
        """Any joining expenses reimbursed by the Company will be recovered in the event you leave the organization within one year of joining.""",
        style='List Number')
    document.add_paragraph(
        """You will be subject to all rules, regulations and policies of the Company, which may be in force from time to time.""",
        style='List Number')
    document.add_paragraph(
        """Any joining expenses / relocation expenses / Notice Pay buyout reimbursed or any other payments made to you while joining by the Company will be recovered in the event you leave the organization within one year of joining.""",
        style='List Number')
    document.add_paragraph("""Please return the duplicate of this letter, duly signed, in token of your acceptance of the above mentioned terms and conditions of the employment, having read the attached Code of Conduct for Senior Management & Employees and on joining you will abide by its prescriptive principles.


	We wish you a long and fruitful career with us. 

	With Regards,

	Yours Sincerely,""")
    sign = document.add_paragraph('')
    sign.add_run("""For MAHINDRA INTEGRATED BUSINESS SOLUTIONS PRIVATE LIMITED





	Riten Chakrabarty
	Chief Finance Officer.""").bold = True

    document.save(path)


def Contract_MS0_MS4(record, path):
    doj = record[58]
    doj = doj[3:5] + '/' + doj[0:2] + '/' + doj[-4:]

    dob = record[28]
    dob = dob[3:5] + '/' + dob[0:2] + '/' + dob[-4:]

    doc = record[62]
    doc = doc[3:5] + '/' + doc[0:2] + '/' + doc[-4:]

    document = Document()

    p = document.add_paragraph(F"""





	{doj}

	Employee Code: {record[57]}

	{record[1]} {record[2]} {record[3]} {record[4]}
	{record[12]},
	{record[13]},{record[15]}

	Dear {record[2]},

	This has reference to the discussions we had regarding taking up employment with us as an {record[67]} on a Fixed Term Contract basis. The terms of the same are enumerated here under:""")

    document.add_paragraph("""The remuneration will be as enumerated in Annexure attached hereto.""",
                           style='List Number')
    a = document.add_paragraph("""This contract will be for a period of """, style='List Number')
    a.add_run('12').bold = True
    a.add_run(F""" Months from """)
    a.add_run(F"""{doj} to {doc}. """).bold = True
    a.add_run(F"""The Contract may be extended based on mutual agreement at the end of this period.""")

    a = document.add_paragraph("""You will not be eligible for any """, style='List Number')
    a.add_run('other benefit / ').bold = True
    a.add_run(F"""facility which is """)
    a.add_run(F"""not mentioned """).bold = True
    a.add_run(F"""in this letter.""")

    document.add_paragraph("""All payments to you shall be subject to deduction of tax at source.""",
                           style='List Number')
    document.add_paragraph(
        """In case you are required to travel outside the city limits, you will be governed by the rules as applicable to an employee of a similar position in the Company equivalent to grade (Grade).""",
        style='List Number')
    document.add_paragraph(
        """The Company also reserves its right to depute you to any of the establishments or its subsidiary.""",
        style='List Number')
    document.add_paragraph(
        """During the Contract period you will be entitled to (Manually Filled) days of compensated absence. Any absenteeism beyond this limit will result in pro-rata deduction of the salary mentioned in (1) above. No encashment of compensated absence will be permissible.""",
        style='List Number')
    document.add_paragraph(
        """It is expected that you will honestly, diligently and efficiently discharge your duties under this contract. In the event of breach of the terms of this contract or substantial misconduct, gross negligence or substantial failure to perform, on your part, during the Contract Period, the Company will be entitled to terminate this contract with immediate effect without any compensation to you. Our decision on the above issues will be final and binding on you.""",
        style='List Number')
    document.add_paragraph(
        """During the Contract Period, you will not engage yourself directly or indirectly in other business or work part-time or accept any other form of employment or contract assignments except to the extent as may be permitted by us in writing. Similarly, you will not engage yourself in any business activity which, in our opinion, is likely to have conflict of interest with us.""",
        style='List Number')
    document.add_paragraph(
        """The Company will be entitled to terminate the contract by giving 15 days’ notice or payment of 15 days remuneration in lieu of the notice. The Company will also be at liberty to pay you 30 days notice wages in lieu of notice. However, it will be impermissible for you to waive the shortfall in the notice period by buying the said shortfall period in lieu thereof except with written permission.""",
        style='List Number')
    document.add_paragraph(
        """During the Contract Period, you will, at all times, observe secrecy and confidentiality and will not divulge, disclose or make known to any unauthorized person within or outside the Company, nor will you use, any knowledge or information in respect of manufacturing, technical trade or business data (including manufacturing processes, technical know-how, customer information, business plans and like matters) which are necessarily confidential and have come to your knowledge and possession.""",
        style='List Number')
    document.add_paragraph(
        """You will also not remove any such information in any form whatsoever from the Company premises, nor copy or transmit the same and nor will you grant permission to assist, permit entry to, or in any manner co-operate with any unauthorized person, for the purposes of accessing, obtaining, copying, transmitting or removing the above.""",
        style='List Number')
    document.add_paragraph(
        """Even after the expiry of your contract with the Company, you will not use, divulge, disclose or remove in any manner whatsoever, confidential information of the type described above of which you were in possession whilst in service, to the detriment of the Company.""",
        style='List Number')
    document.add_paragraph(
        """You will also observe all the confidentiality measures which are in existence, or which may be enforced from time to time, as well as directions as to confidentiality marked on any communication, document, computer floppy etc.""",
        style='List Number')
    document.add_paragraph(
        """You shall indemnify and hold Company harmless and indemnified against any damage or loss caused to the Company on account of breach of confidentiality on your part.""",
        style='List Number')
    document.add_paragraph(
        """These confidentiality provisions shall survive the termination / expiry of your employment with the Company, either by way of termination of this contract or otherwise.""",
        style='List Number')
    document.add_paragraph(
        """You will assign to the Company your entire right, title and interest in any Intellectual Property Rights (IPRs for short, which term would include patents, trade-marks, copyrights, designs, whether registered or not, and all improvements thereto) that you may make, solely or jointly with others, in the course of the Contract with the Company relating to any or all systems, services and products manufactured or marketed or leased or developed.""",
        style='List Number')
    document.add_paragraph(
        """You will perform all necessary acts and execute such documents in such format as may be required by the Company, without expense to you, which in the judgment of the Company or its Attorneys may be necessary or desirable to secure to the Company full right, title and interest in the IPRs.""",
        style='List Number')
    document.add_paragraph(
        """This contract does not in any respect make either of us, an agent or a partner of the other or authorize either to transact any business in the name of the other or to incur any obligation or liability for or on behalf of each other, unless specifically authorized by a special power of attorney by either of us, duly executed under the relevant laws.""",
        style='List Number')
    document.add_paragraph(
        """The Company shall at all times have the right to access and monitor all e-mails created, sent / received or stored by you, using Company facility and on Company's system, at any time, without any prior notification. All such data and information shall belong and continue to remain the property of the Company at all times.""",
        style='List Number')
    document.add_paragraph(
        """Even after the expiry of your contract with the Company, for whatever reasons, you shall forthwith, return all company properties, movable and immovable, including, without limitation, all Company information, files, reports, memoranda, software, credit cards, door and file keys, computer access codes and such other property which you had directly or indirectly was privy to or received or is / was in possession, or prepared during your association with the Company.""",
        style='List Number')
    document.add_paragraph(
        """You shall at all times endeavor to protect Proprietary and Confidential information, electronic information, electronic form, electronic record, information within your purview, systems & processes from unauthorized disclosure, alteration or destruction.""",
        style='List Number')
    document.add_paragraph(
        """This contract constitutes a fixed term contract of employment and you shall not transfer or assign this contract either wholly or in part, or any rights and obligations hereunder. We shall be entitled to assign this contract to any of our subsidiaries, affiliates or associates.""",
        style='List Number')
    document.add_paragraph("""This contract is subject to the jurisdiction of the Courts of Mumbai.""",
                           style='List Number')
    document.add_paragraph("""This letter is being issued in duplicate. If these terms and conditions are acceptable to you, please return the duplicate of this letter duly counter signed by you in token of your acceptance hereof.

	Yours Sincerely,
	""")
    sign = document.add_paragraph('')
    sign.add_run("""For MAHINDRA INTEGRATED BUSINESS SOLUTIONS PRIVATE LIMITED





	Riten Chakrabarty
	Chief Finance Officer.""").bold = True

    document.save(path)


def Trainee(record, path):
    doj = record[58]
    doj = doj[3:5] + '/' + doj[0:2] + '/' + doj[-4:]

    dob = record[28]
    dob = dob[3:5] + '/' + dob[0:2] + '/' + dob[-4:]

    document = Document()

    p = document.add_paragraph(F"""





	{doj}

	Employee Code: {record[57]}

	{record[1]} {record[2]} {record[3]} {record[4]}
	{record[12]},
	{record[13]},{record[15]}

	Dear {record[2]},

	We have pleasure of offering you appointment in our company, as {record[67]} – {record[84]} in Band {record[66]}.

	This offer of appointment is subject to you satisfying the following:""")
    document.add_paragraph(
        F"""Your written undertaking to join, not later than {doj}. You have been deputed to work at {record[70]}.""",
        style='List Bullet')
    document.add_paragraph(
        F"""Submission of all your necessary certificates and documents in respect of educational and professional qualifications, proof of age and previous employment, as per the requirements of the Company.""",
        style='List Bullet')

    document.add_paragraph(F"""This appointment will commence from the date on which you actually join the Company. 

	The terms and conditions of employment shall be as under:""")

    document.add_paragraph("""You will receive emoluments / allowances as per the attached Annexure.""",
                           style='List Number')
    a = document.add_paragraph("""Your employment will be on the rolls of """, style='List Number')
    a.add_run("'Mahindra Integrated Business Solutions Pvt. Ltd.'").bold = True

    document.add_paragraph(
        """Gratuity, Provident Fund, and Medical Benefit and family, as per the rules of the Company, Medical Benefit will be applicable from the date of confirmation in service.""",
        style='List Number')

    document.add_paragraph(
        F"""The age of retirement will be sixty years. (On the strength of the bio-data submitted by you, we have recorded your date of birth as {dob})""",
        style='List Number')
    document.add_paragraph(
        """With effect from the date of your employment, you are required to become a member Provident Fund.""",
        style='List Number')
    document.add_paragraph(
        """You are liable to be placed for service at our discretion at any of the Company's establishments/departments/divisions anywhere in India as also the Associate Companies and subsidiary Companies of Mahindra & Mahindra Ltd.""",
        style='List Number')
    document.add_paragraph(
        """You will be on probation for a period of 06 months. This probationary period could be curtailed or extended at the discretion of the Company. However, completion of 06 months of probation does not entitle you or result in automatic confirmation of your employment, unless the Company confirms your employment in writing. During this period, your employment may be terminated forthwith without notice and without assigning any reasons.""",
        style='List Number')
    document.add_paragraph(
        """You will be entitled to Exigency Leave as per the rules. Upon satisfactory completion of the period of your probation, you will be confirmed in our service and on confirmation:""",
        style='List Number')
    document.add_paragraph("""You will be entitled to Privilege Leave, as per the rules thereof;""",
                           style='List Bullet')
    document.add_paragraph(
        """The Company will be at liberty to terminate your services with 90 days notice or by paying you 90 days salary, including allowances, in lieu of notice. In the event the Company decides to pay you 90 days salary in lieu of notice, the Company will be at liberty to call upon you not to take up any alternate employment for the period of 90 days. The Company will also be at liberty to call upon you not to report for work, though you would be on the rolls of the Company for the said period and you would be paid your salary as per your contract, as if you were on duty. In the event you choose to resign from the services of the Company, you will be required to serve for the period of notice of 90 days. The Company, however, will be at liberty to call upon you not to report for work or even take up any alternate employment during this period, which will be at the sole discretion of the Company. The Company will also be at liberty to pay you 90 days notice wages in lieu of notice. However, it will be impermissible for you to waive the shortfall in the notice period by buying the said shortfall period in lieu thereof except with written permission.""",
        style='List Bullet')
    document.add_paragraph(
        """During the probationary period, the notice period to be served by either party shall be 15 days.""",
        style='List Bullet')
    document.add_paragraph(
        """Further, you shall not be entitled to adjust your notice period against privilege leave, if any, standing to you credit.""",
        style='List Bullet')
    document.add_paragraph(
        """So long as you are in the employment of the Company, you will, at all times, observe secrecy and confidentiality and will not divulge, disclose or make known to any unauthorized person within or outside the Company, nor will you unauthorized use any knowledge or information in respect of manufacturing, technical trade or business data (including manufacturing processes, technical know-how, customer information, business plans and like matters) which are necessarily confidential and have come to your knowledge and possession.  You will also not remove any such information in any form whatsoever from the Company premises, nor copy or transmit the same unauthorized nor will you grant permission to assist, permit entry to, or in any manner co-operate with any unauthorized person for the purposes of accessing, obtaining, copying, transmitting or removing the above. Even after the cessation of your employment with the Company, you will not use, divulge, disclose or remove in any manner whatsoever confidential information of the type described above of which you were in possession whilst in service to the detriment of the Company.   You will also observe all the confidentiality measures which are in existence, or which may be enforced from time to time, as well as directions as to confidentiality marked on any communication, document, computer floppy etc. You shall indemnify and hold Company harmless and indemnified against any damage or loss caused to the Company on account of breach of confidentiality on your part. These confidentiality provisions shall survive the separation of your employment with the Company, either by way of retirement or termination or otherwise.""",
        style='List Number')
    document.add_paragraph(
        """In addition to your fulfilling the requirements of secrecy and confidentiality, as specified herein, also during your employment with the Company, you shall not engage in any vocation, training, employment, consultancy, business, transaction, or any other activity, which is in conflict with the interests of the Company, in any capacity whatsoever either on your own or in association with any other individual/firm/institute/body corporate, etc., whether for any consideration or not.""",
        style='List Number')
    document.add_paragraph(
        """You will devote your full attention exclusively to the duties entrusted to you from time to time by the Company and while in service of this Company you will not work for any person or Company in any capacity either for any consideration or otherwise, nor do any private business without obtaining prior permission of the Company in writing.""",
        style='List Number')
    document.add_paragraph(
        """You will assign to the Company your entire right, title and interest in any Intellectual Property Rights (IPRs for short, which term would include patents, trade-marks, copyrights, designs, whether registered or not, and all improvements thereto) that you may make, solely or jointly with others, in the course of your employment with the Company relating to any or all systems, services and products manufactured or marketed or leased or developed.  You  will perform all necessary  acts and  execute such documents in such format as may be required by the Company,  without  expense  to you, which in the  judgment  of  the  Company  or its Attorneys may be necessary or desirable to  secure  to the Company full right title and interest in the IPRs.""",
        style='List Number')
    document.add_paragraph(
        """The Company shall at all times have the right to access and monitor all e-mails created, sent / received or stored by you using Company facility and on Company’s system at any time without giving you any prior notification. All such data and information shall be the property of the Company at all times.""",
        style='List Number')
    document.add_paragraph(
        """You shall endeavor to uphold the good image of the Company and shall not by your conduct adversely affect the reputation of the Company and bring disrepute to the Company, in any manner whatsoever.""",
        style='List Number')
    document.add_paragraph(
        """You shall, on ceasing to be the employee of the Company, forthwith return all Company properties, movable and immovable, including, without limitation, all Company information, files, reports, memoranda, software, credit cards, door and file keys, computer access codes and such other property which you received or in possession or prepared in connection with your employment with the Company""",
        style='List Number')
    document.add_paragraph(
        """Any joining expenses reimbursed by the Company will be recovered in the event you leave the organization within one year of joining.""",
        style='List Number')
    document.add_paragraph(
        """You will be subject to all rules, regulations and policies of the Company, which may be in force from time to time.""",
        style='List Number')
    document.add_paragraph(
        """Any joining expenses / relocation expenses / Notice Pay buyout reimbursed or any other payments made to you while joining by the Company will be recovered in the event you leave the organization within one year of joining.""",
        style='List Number')
    document.add_paragraph("""Please return the duplicate of this letter, duly signed, in token of your acceptance of the above mentioned terms and conditions of the employment, having read the attached Code of Conduct for Senior Management & Employees and on joining you will abide by its prescriptive principles.


	We wish you a long and fruitful career with us. 

	With Regards,

	Yours Sincerely,""")
    sign = document.add_paragraph('')
    sign.add_run("""For MAHINDRA INTEGRATED BUSINESS SOLUTIONS PRIVATE LIMITED





	Riten Chakrabarty
	Chief Finance Officer.""").bold = True

    document.save('demo.docx')


def get_letter(record, key, path):
    if key == "On_Contract":
        On_Contract(record, path)
    if key == "On_rolls":
        On_rolls(record, path)
    if key == "L_Band":
        L_Band(record, path)
    if key == "MB_and_Above":
        MB_and_Above(record, path)
    if key == "MT":
        MT(record, path)
    if key == "MS0_MS4":
        MS0_MS4(record, path)
    if key == "Contract_MS0_MS4":
        Contract_MS0_MS4(record, path)
    if key == "Trainee":
        Trainee(record, path)
